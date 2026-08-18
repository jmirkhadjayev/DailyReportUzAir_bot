"""Hisobotlarni Word (.docx) formatiga chiqarish."""
from __future__ import annotations

from collections import OrderedDict
from datetime import date
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config import FOOTER_TEXT, ORG_NAME, chief_name
from i18n import t
from utils import fmt_date_long, now_str, period_title

ACCENT = RGBColor(0x1F, 0x4E, 0x78)
COLUMN_KEYS = ["col_num", "col_tabel", "col_name_position", "col_done", "col_problems", "col_plans"]
COL_WIDTHS = [Cm(0.9), Cm(1.6), Cm(4.8), Cm(8.6), Cm(5.4), Cm(5.4)]


def _shade(cell, color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shd)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _set_cell(cell, text: str, *, bold: bool = False, size: int = 10, center: bool = False,
              color: RGBColor | None = None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(2)
    run = paragraph.add_run()
    for i, line in enumerate((text or "—").split("\n")):
        if i:
            run.add_break()
        run.add_text(line)
    run.bold = bold
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def _setup(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    if section.page_width < section.page_height:
        section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.5)
    section.top_margin = section.bottom_margin = Cm(1.5)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # har bir sahifaning eng pastida chiqadigan imzo
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(FOOTER_TEXT)
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def build_word(
    rows, date_from: date, date_to: date, lang: str = "uz", subject: str = ""
) -> bytes:
    document = Document()
    _setup(document)
    columns = [t(lang, key) for key in COLUMN_KEYS]

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(ORG_NAME.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(t(lang, "doc_title"))
    run.bold = True
    run.font.size = Pt(13)

    if subject:
        line = document.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run(subject)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = ACCENT

    info = document.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(
        f"{t(lang, 'doc_period')}: {period_title(date_from, date_to, lang)}\n"
        f"{t(lang, 'doc_generated')}: {now_str()}   |   "
        f"{t(lang, 'doc_total')}: {len(rows)}"
    )
    run.font.size = Pt(10)
    run.italic = True

    if not rows:
        empty = document.add_paragraph()
        empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        empty.add_run("\n" + t(lang, "doc_empty")).bold = True
        return _to_bytes(document)

    by_date: "OrderedDict[str, list]" = OrderedDict()
    for r in rows:
        by_date.setdefault(r["report_date"], []).append(r)

    for index, (day, day_rows) in enumerate(by_date.items()):
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(14 if index else 10)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run(
            f"{fmt_date_long(day, lang)}  —  {t(lang, 'doc_reports_count', count=len(day_rows))}"
        )
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = ACCENT

        table = document.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        header_cells = table.rows[0].cells
        for i, name in enumerate(columns):
            _set_cell(header_cells[i], name, bold=True, size=10, center=True,
                      color=RGBColor(0xFF, 0xFF, 0xFF))
            _shade(header_cells[i], "1F4E78")
        _repeat_header(table.rows[0])

        for i, r in enumerate(day_rows, start=1):
            cells = table.add_row().cells
            _set_cell(cells[0], str(i), center=True)
            _set_cell(cells[1], r["tabel"] or "—", center=True)
            _set_cell(cells[2], f"{r['full_name']}\n{r['position'] or '—'}")
            _set_cell(cells[3], r["done"])
            _set_cell(cells[4], r["problems"])
            _set_cell(cells[5], r["plans"])
            if i % 2 == 0:
                for cell in cells:
                    _shade(cell, "F2F7FC")

        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = COL_WIDTHS[i]

    _signature_block(document, lang)
    return _to_bytes(document)


def _signature_block(document: Document, lang: str) -> None:
    document.add_paragraph()
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(24)
    run = paragraph.add_run(t(lang, "doc_sign", name=chief_name(lang)))
    run.font.size = Pt(11)
    run.bold = True
    note = document.add_paragraph()
    note.add_run(t(lang, "doc_sign_note")).font.size = Pt(9)


def _to_bytes(document: Document) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
